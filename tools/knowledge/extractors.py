"""Format-specific text extractors for the TouriBot knowledge ingest pipeline.

Each extractor receives a file Path and returns clean markdown text,
or an empty string if the file cannot be processed.
"""

import re
from pathlib import Path


def extract_html(path: Path) -> str:
    """Extract markdown from a wiki HTML file with window.__MD_EN template literal."""
    try:
        html = path.read_text(encoding="utf-8", errors="replace")
        match = re.search(r'window\.__MD_EN\s*=\s*`(.*?)`\s*;', html, re.DOTALL)
        if not match:
            return ""
        md = match.group(1).strip()
        return _clean_wiki_markdown(md)
    except Exception:
        return ""


def extract_txt(path: Path) -> str:
    """Extract content from a plain-text file."""
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
        return text.strip()
    except Exception:
        return ""


def extract_md(path: Path) -> str:
    """Read a markdown file as-is, preserving frontmatter."""
    try:
        return path.read_text(encoding="utf-8", errors="replace").strip()
    except Exception:
        return ""


def extract_pdf(path: Path) -> str:
    """Extract text from a PDF file using pdfplumber, with page markers."""
    try:
        import logging
        import pdfplumber
        # Suppress noisy font-descriptor warnings from pdfminer
        logging.getLogger("pdfminer").setLevel(logging.ERROR)
        pages = []
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                text = text.strip()
                if text:
                    pages.append(f"<!-- Page {i} -->\n{text}")
        return "\n\n".join(pages)
    except ImportError:
        return ""
    except Exception:
        return ""


def extract_docx(path: Path) -> str:
    """Extract paragraphs and table cells from a .docx file."""
    try:
        from docx import Document
        doc = Document(path)
        parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Render headings as markdown
                if para.style and para.style.name.startswith("Heading"):
                    try:
                        level = int(para.style.name.split()[-1])
                    except (ValueError, IndexError):
                        level = 2
                    parts.append("#" * min(level, 6) + " " + text)
                else:
                    parts.append(text)

        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                parts.append("\n".join(rows))

        return "\n\n".join(parts)
    except ImportError:
        return ""
    except Exception:
        return ""


def extract_auto(path: Path) -> str:
    """Dispatch to the correct extractor based on file extension."""
    suffix = path.suffix.lower()
    if suffix == ".html":
        return extract_html(path)
    elif suffix == ".txt":
        return extract_txt(path)
    elif suffix == ".md":
        return extract_md(path)
    elif suffix == ".pdf":
        return extract_pdf(path)
    elif suffix in (".docx", ".doc"):
        return extract_docx(path)
    return ""


# ── Internal helpers ──────────────────────────────────────────────────────────


def _clean_wiki_markdown(md: str) -> str:
    """Clean up markdown extracted from a JS template literal."""
    md = md.replace("\\`", "`")
    md = md.replace("\\$", "$")
    md = re.sub(r'\\u([0-9a-fA-F]{4})', lambda m: chr(int(m.group(1), 16)), md)
    md = md.replace("\u200b", "")
    md = md.replace("\r\n", "\n")
    md = re.sub(r'\n{4,}', '\n\n\n', md)
    return md.strip()
